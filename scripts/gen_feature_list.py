"""Generate 系统功能清单.docx — 全模��功能列表"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_para(text, bold=False, color=None, size=10.5, left_indent=0):
    p = doc.add_paragraph()
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("▶  " + text)
    _set_font(run, size=10, color=(0x64, 0x74, 0x8B))
    return p


def make_table(headers, rows, header_bg="1E40AF", col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_font(run, bold=True, color=(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, "F1F5F9")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_font(run)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def gap():
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
gap()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("海外电商分析系统  功能清单")
_set_font(r, size=22, bold=True, color=(0x1E, 0x40, 0xAF))

gap()
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("全模块功能汇总  v1.0  /  2026-04")
_set_font(r2, size=12, color=(0x64, 0x74, 0x8B))
gap()

# ═══════════════════════════════════════════════════════════
# 一、登录
# ═══════════════════════════════════════════════════════════
add_heading("一、登录", 1)
add_note("路由：/auth/login")
make_table(
    ["功能点", "说明"],
    [
        ["账号/密码输入框", "标准文本输入；密码框可切换明文/密文（眼睛图标）"],
        ["登录按钮", "点击后调用后端验证；加载期间按钮禁用并显示「登录中…」"],
        ["错误提示", "账号不存在、密码错误时显示红色错误信息"],
        ["跳转逻辑", "登录成功后跳转至 callbackUrl（或默认首页 /work）"],
        ["Session 维持", "服务端 Session（NextAuth）；刷新页面保持登录状态"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 二、首页
# ═══════════════════════════════════════════════════════════
add_heading("二、首页", 1)
add_note("路由：/work  或  /work/dashboard/inventory-turnover-board")
make_table(
    ["功能点", "说明"],
    [
        ["库存周转看板", "展示库存周转相关数据（当前为占位/图表页面）"],
        ["导航入口", "通过左侧菜单跳转到各子模块"],
        ["权限保护", "未登录自动跳转 /auth/login；无权限菜单项不可见"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 三、选品
# ═══════════════════════════════════════════════════════════
add_heading("三、选品（ops.selection）", 1)
add_note("路由：/work/ops/selection")

add_heading("3.1  搜索与筛选", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品名称搜索", "关键词模糊搜索"],
        ["所属类目下拉", "按类目过滤"],
        ["时间范围下拉", "今天 / 7日内 / 15日内 / 30日内 / 90日内 / 180日内 / 自定义日期范围（共7种）"],
        ["查询按钮", "触发接口查询；加载中显示「查询中…」"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("3.2  状态卡片", 2)
make_table(
    ["功能点", "说明"],
    [
        ["全部商品", "显示全量记录总数；点击不过滤"],
        ["选品中", "客户端即时过滤，无需重新查询"],
        ["已淘汰", "客户端即时过滤"],
        ["已选品", "客户端即时过滤"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("3.3  历史数据", 2)
make_table(
    ["功能点", "说明"],
    [
        ["查看历史数据按钮", "切换为历史数据视图；操作列按钮隐藏（只读模式）"],
        ["退出历史数据", "返回正常列表；卡片和操作按钮恢复"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("3.4  新建记录", 2)
make_table(
    ["功能点", "说明"],
    [
        ["新建按钮", "打开新建弹窗"],
        ["基本信息字段", "商品名称、所属类目、产品规格、TEMU商品链接、平台在售价格（Min/Max）"],
        ["产品尺寸", "长/宽/高（cm）"],
        ["产品重量", "kg"],
        ["单套尺寸", "长/宽/高（cm）"],
        ["包裹实重", "kg"],
        ["产品图片", "图片上传（预览+删除）"],
        ["参考链接", "URL 输入"],
        ["批量规格", "可新增多行规格，每行含颜色/尺码/SKU等子字段"],
        ["提交创建", "保存并进入选品流程"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("3.5  表格列", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品基本信息", "图片缩略图（可点击预览）+ 商品名称 + ID + 类目 + 参考链接"],
        ["产品属性", "产品尺寸（长x宽x高 cm）+ 重量（kg）+ 规格"],
        ["单套属性", "单套尺寸 + 包裹实重"],
        ["平台在售价格", "Min - Max 区间"],
        ["状态", "彩色圆点 + 状态文字"],
        ["操作", "修改 / 淘汰 / 转询价 按钮（按状态显示）"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("3.6  编辑弹窗", 2)
make_table(
    ["功能点", "说明"],
    [
        ["所有字段可编辑", "与新建弹窗字段相同"],
        ["单位切换", "产品属性/单套属性支持 cmkg ↔ 英寸/英镑 切换显示（不写入DB）"],
        ["保存", "保存修改，状态不变"],
        ["淘汰", "状态变为已淘汰"],
        ["转询价", "记录进入询价流程（状态变更）"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 四、询价
# ═══════════════════════════════════════════════════════════
add_heading("四、询价（ops.inquiry）", 1)
add_note("路由：/work/ops/inquiry")

add_heading("4.1  搜索与筛选", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品名称搜索", "关键词模糊搜索"],
        ["所属类目下拉", "按类目过滤"],
        ["时间范围下拉", "同选品，含7种选项"],
        ["负责人筛选", "按指定的询价负责人过滤"],
        ["查询按钮", "触发接口查询"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("4.2  状态卡片", 2)
make_table(
    ["功能点", "说明"],
    [
        ["全部询价商品", "全量记录"],
        ["待询价", "客户端即时过滤"],
        ["询价中", "客户端即时过滤"],
        ["已询价", "客户端即时过滤"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("4.3  历史数据", 2)
make_table(
    ["功能点", "说明"],
    [
        ["查看/退出历史数据", "同选品，只读模式"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("4.4  批量操作", 2)
make_table(
    ["功能点", "说明"],
    [
        ["批量分配负责人", "勾选多条记录后指定询价负责人"],
        ["批量编辑", "同时修改多条记录的某些字段"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("4.5  编辑弹窗", 2)
make_table(
    ["功能点", "说明"],
    [
        ["基本信息（只读）", "商品名称、所属类目、产品图片、参考链接、平台在售价格（只读）"],
        ["产品属性（只读+单位切换）", "产品尺寸长/宽/高 + 重量 + 规格；支持 cmkg ↔ 英寸/英镑"],
        ["单套属性（只读+单位切换）", "单套尺寸 + 包裹实重；支持单位切换"],
        ["询价信息", "询价负责人（可编辑）、询价状态"],
        ["成本信息", "采购成本（RMB）、头程成本（RMB）、美元汇率"],
        ["自动计算预览", "体积重、尾程成本、temu报价 实时计算显示"],
        ["保存", "保存修改，状态不变"],
        ["提交转核价", "状态变为待核价"],
        ["撤回弹窗", "需填写撤回理由（必填）；确认后状态回退"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("4.6  表格列", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品基本信息", "图片 + 名称 + ID + 类目 + 参考链接"],
        ["产品属性", "尺寸 + 重量 + 规格（含 cmkg/英寸英镑 切换）"],
        ["单套属性", "尺寸 + 重量"],
        ["成本总计（RMB）", "大号金额 + 明细（采购/头程/尾程/负向）"],
        ["TEUM供货价", "自动计算值"],
        ["参考销售售价", "Min - Max 区间"],
        ["负责人", "当前询价负责人"],
        ["状态", "彩色圆点 + 文字"],
        ["操作", "修改 / 撤回 / 转核价"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 五、核价
# ═══════════════════════════════════════════════════════════
add_heading("五、核价（ops.pricing）", 1)
add_note("路由：/work/ops/pricing")

add_heading("5.1  搜索与筛选", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品名称搜索", "关键词模糊搜索"],
        ["所属类目下拉", "按类目过滤"],
        ["时间范围下拉", "7种时间选项"],
        ["负责人筛选", "按核价负责人过滤"],
        ["查询按钮", "触发接口查询"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("5.2  状态卡片", 2)
make_table(
    ["功能点", "说明"],
    [
        ["全部核价商品", "全量记录"],
        ["待核价", "客户端即时过滤"],
        ["核价中", "客户端即时过滤"],
        ["待确品", "客户端即时过滤"],
        ["已废弃", "客户端即时过滤"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("5.3  批量操作", 2)
make_table(
    ["功能点", "说明"],
    [
        ["批量分配负责人", "勾选多条后指定核价负责人"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("5.4  历史数据", 2)
make_table(
    ["功能点", "说明"],
    [
        ["查看/退出历史数据", "只读模式，操作按钮隐藏"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("5.5  编辑弹窗", 2)
make_table(
    ["功能点", "说明"],
    [
        ["基本信息（只读）", "同询价，含图片/链接/平台在售价格"],
        ["产品属性（只读+单位切换）", "长/宽/高/重量；cmkg ↔ 英寸/英镑"],
        ["单套属性（只读+单位切换）", "单套尺寸 + 包裹实重"],
        ["核价信息", "采购成本（可编辑）、头程成本（可编辑）、美元汇率（可编辑）"],
        ["只读计算字段", "尾程成本（RMB）、负向成本、TEUM供货价"],
        ["成本总计", "大号实时合计，含明细"],
        ["参考销售售价", "只读"],
        ["保存", "保存修改，状态不变"],
        ["提交转确品", "状态变为待确品"],
        ["废弃弹窗", "填写废弃理由（必填），确认后状态变为已废弃"],
        ["撤回弹窗", "填写撤回理由（必填），回退至询价"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("5.6  表格列", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品基本信息", "图片 + 名称 + ID + 类目 + 链接"],
        ["产品属性", "尺寸 + 重量 + 规格（列头含 cmkg/英寸英镑 切换按钮）"],
        ["单套属性", "尺寸 + 重量"],
        ["成本总计（RMB）", "金额 + 明细"],
        ["TEUM供货价", "自动计算值"],
        ["参考销售售价", "Min - Max"],
        ["负责人", "当前核价负责人"],
        ["状态", "彩色圆点 + 文字"],
        ["操作", "按状态不同显示：修改 / 提交 / 废弃 / 撤回"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 六、确品
# ═══════════════════════════════════════════════════════════
add_heading("六、确品（ops.confirm）", 1)
add_note("路由：/work/ops/confirm  |  数据存储在 ops.purchase 表")

add_heading("6.1  搜索与筛选", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品名称搜索", "关键词模糊搜索"],
        ["所属类目下拉", "按类目过滤"],
        ["时间范围下拉", "7种时间选项"],
        ["查询按钮", "触发接口查询；加载中显示「查询中…」"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("6.2  状态卡片", 2)
make_table(
    ["功能点", "说明"],
    [
        ["全部待确品商品", "待确品 + 待采购 记录合计"],
        ["待确品", "客户端即时过滤"],
        ["待采购", "客户端即时过滤"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("6.3  历史数据", 2)
make_table(
    ["功能点", "说明"],
    [
        ["查看历史数据按钮", "切换历史视图；状态卡片半透明；操作按钮隐藏"],
        ["退出历史数据", "恢复正常模式"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("6.4  表格列（16列）", 2)
make_table(
    ["列名", "说明"],
    [
        ["商品基本信息", "图片缩略图（点击预览）+ 名称 + ID + 类目（ID | 类目名格式）+ 参考链接"],
        ["产品属性", "尺寸（长x宽x高 cm）+ 重量 + 规格"],
        ["单套属性", "单套尺寸 + 包裹实重"],
        ["成本总计（RMB）", "大号金额 + 明细（采购/头程/尾程/负向）"],
        ["TEUM供货价", "temu报价对应金额"],
        ["参考销售售价", "Min - Max 区间"],
        ["SKU信息", "公司商品编码 + 仓库条码 + 平台编码"],
        ["发货信息", "各仓库发货量"],
        ["销量预估", "预估销量（大号）+ 预估来源（小字）"],
        ["实际下单数", "实际采购下单数量"],
        ["订单总数", "累计订单数"],
        ["交货周期", "交货天数"],
        ["发货日期", "计划发货日期"],
        ["发货数量", "本次发货数量"],
        ["状态", "橙色圆点（待确品）/ 蓝色圆点（待采购）"],
        ["操作", "修改 + 撤回（待采购状态才可撤回）"],
    ],
    col_widths=[4, 12],
)
gap()

add_heading("6.5  编辑弹窗（7个区块）", 2)
make_table(
    ["区块", "字段与功能"],
    [
        ["基本信息（只读）", "产品图片（点击大图预览）/ 名称 / 所属类目 / 参考链接（新标签打开）/ 参考销售售价（Min/Max只读）"],
        ["产品属性（只读+单位切换）", "长/宽/高（cm） / 产品重量 / 产品规格；支持 cmkg ↔ 英寸/英镑 实时换算（不写入DB）"],
        ["单套属性（只读+单位切换）", "单套长/宽/高 / 包裹实重；支持单位切换"],
        ["核价信息", "成本总计大号展示 / 采购成本（可编辑）/ 头程成本（可编辑）/ 尾程成本（只读）/ 负向成本（只读）/ TEUM供货价（只读）/ 参考销售售价（只读）"],
        ["SKU信息", "公司商品编码（可编辑）/ 仓库条码（多行，可增删，美西仓/美东仓等）/ 平台编码（多行，可增删，TEUM等）"],
        ["发货信息", "多行仓库发货量（可增删，仓库名+数量）"],
        ["销售预估", "预估来源（可编辑）/ 预估销量（可编辑）"],
    ],
    col_widths=[4, 12],
)
gap()

add_heading("6.6  弹窗操作按钮", 2)
make_table(
    ["功能点", "说明"],
    [
        ["保存", "写入修改，状态维持不变（待确品）"],
        ["提交", "状态变为待采购"],
        ["取消", "关闭弹窗，不保存"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("6.7  撤回弹窗", 2)
make_table(
    ["功能点", "说明"],
    [
        ["触发条件", "状态为待采购的记录，点「撤回」按钮"],
        ["弹窗内容", "商品信息（只读）+ 撤回理由（必填文本框）"],
        ["必填验证", "未填写理由时「确认撤回」按钮禁用"],
        ["确认撤回", "状态回退到待核价；记录在确品页消失，核价页出现"],
        ["取消", "弹窗关闭，状态不变"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("6.8  输入交互细节", 2)
make_table(
    ["功能点", "说明"],
    [
        ["数字输入框禁止负数", "采购成本/头程成本等框不允许输入负号"],
        ["禁止科学计数法", "e 字符被过滤"],
        ["禁止滚轮改值", "wheel 事件已阻止"],
        ["中文全角标点转换", "输入「，」自动转「.」；「。」自动转「.」（用于价格输入）"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 七、采购
# ═══════════════════════════════════════════════════════════
add_heading("七、采购（ops.purchase）", 1)
add_note("路由：/work/ops/purchase")

add_heading("7.1  搜索与筛选", 2)
make_table(
    ["功能点", "说明"],
    [
        ["商品名称搜索", "关键词模糊搜索"],
        ["所属类目下拉", "按类目过滤"],
        ["时间范围下拉", "7种时间选项"],
        ["查询按钮", "触发接口查询"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("7.2  状态卡片", 2)
make_table(
    ["功能点", "说明"],
    [
        ["全部采购商品", "全量记录"],
        ["待采购", "客户端即时过滤"],
        ["采购中", "客户端即时过滤"],
        ["已采购", "客户端即时过滤"],
        ["已取消", "客户端即时过滤"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("7.3  历史数据", 2)
make_table(
    ["功能点", "说明"],
    [
        ["查看/退出历史数据", "只读模式"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("7.4  表格列（13列）", 2)
make_table(
    ["列名", "说明"],
    [
        ["商品基本信息", "图片 + 名称 + ID + 类目 + 链接"],
        ["产品属性", "尺寸 + 重量 + 规格"],
        ["单套属性", "单套尺寸 + 重量"],
        ["成本总计（RMB）", "金额 + 明细"],
        ["TEUM供货价", "供货价"],
        ["参考销售售价", "Min - Max"],
        ["SKU信息", "公司编码 + 仓库条码 + 平台编码"],
        ["发货信息", "仓库发货量"],
        ["销量预估", "预估量 + 来源"],
        ["实际下单数", "采购下单数"],
        ["订单总数", "累计订单"],
        ["状态", "彩色圆点 + 文字"],
        ["操作", "修改 / 撤回"],
    ],
    col_widths=[4, 12],
)
gap()

add_heading("7.5  编辑弹窗", 2)
make_table(
    ["功能点", "说明"],
    [
        ["基本信息（只读）", "同确品弹窗"],
        ["产品属性/单套属性", "只读 + 单位切换"],
        ["核价信息", "采购/头程成本（可编辑）+ 只读计算字段"],
        ["SKU信息", "编码与条码（可编辑，可增删）"],
        ["发货信息", "仓库发货量（可增删）"],
        ["采购信息", "实际下单数（可编辑）/ 交货周期 / 发货日期 / 发货数量"],
        ["销售预估", "预估来源 + 预估销量（可编辑）"],
        ["保存", "状态不变"],
        ["提交", "状态推进"],
        ["撤回", "填理由，回退状态"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 八、财务
# ═══════════════════════════════════════════════════════════
add_heading("八、财务模块", 1)

add_heading("8.1  销售数据（finance.sales-data）", 2)
add_note("路由：/work/finance/sales-data")
make_table(
    ["功能点", "说明"],
    [
        ["搜索与筛选", "商品名称搜索 + 时间范围下拉 + 查询按钮"],
        ["表格", "动态 Schema，根据配置展示列"],
        ["编辑弹窗", "可编辑销售数据字段，保存写入"],
        ["历史数据", "查看/退出历史数据（只读模式）"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("8.2  仓库成本（finance.warehouse-cost）", 2)
add_note("路由：/work/finance/warehouse-cost")
make_table(
    ["功能点", "说明"],
    [
        ["搜索与筛选", "商品名称搜索 + 时间范围下拉 + 查询按钮"],
        ["表格", "展示仓库成本明细"],
        ["编辑弹窗", "可编辑成本字段"],
        ["历史数据", "只读历史视图"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("8.3  罚款金额（finance.penalty-amount）", 2)
add_note("路由：/work/finance/penalty-amount")
make_table(
    ["功能点", "说明"],
    [
        ["搜索与筛选", "商品名称搜索 + 时间范围下拉 + 查询按钮"],
        ["表格", "展示罚款金额记录"],
        ["编辑弹窗", "可编辑罚款相关字段"],
        ["历史数据", "只读历史视图"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 九、配置管理
# ═══════════════════════════════════════════════════════════
add_heading("九、配置管理", 1)

add_heading("9.1  用户管理（settings/users）", 2)
add_note("路由：/work/settings/users")
make_table(
    ["功能点", "说明"],
    [
        ["用户列表", "展示所有用户（用户名、邮箱、角色、状态、创建时间）"],
        ["搜索", "关键词搜索用户名/邮箱"],
        ["新建用户", "填写用户名、邮箱、密码、角色，提交创建"],
        ["编辑用户", "修改用户名、邮箱、角色"],
        ["禁用/启用用户", "切换用户账号状态"],
        ["重置密码", "管理员为用户设置新密码"],
        ["删除用户", "永久删除（有二次确认）"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("9.2  角色管理（settings/roles）", 2)
add_note("路由：/work/settings/roles")
make_table(
    ["功能点", "说明"],
    [
        ["角色列表", "展示所有角色（角色名、描述、用户数、创建时间）"],
        ["新建角色", "填写角色名称和描述"],
        ["编辑角色", "修改角色名称和描述"],
        ["删除角色", "有二次确认；关联用户时不允许删除"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("9.3  类目配置（settings/categories）", 2)
add_note("路由：/work/settings/categories")
make_table(
    ["功能点", "说明"],
    [
        ["类目列表", "展示所有类目名称及创建时间"],
        ["新建类目", "输入类目名称，提交保存"],
        ["编辑类目", "修改类目名称"],
        ["删除类目（二步确认）", "Step1：点击删除按钮，按钮变为红色「确认删除？」；Step2：再次点击才真正删除；1秒内未二次点击自动还原"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("9.4  尾程价目表（settings/last-mile-pricing）", 2)
add_note("路由：/work/settings/last-mile-pricing")
make_table(
    ["功能点", "说明"],
    [
        ["价目表列表", "展示各仓库/重量段的尾程运费配置"],
        ["搜索", "按仓库名称或条件关键词搜索"],
        ["新建价目", "填写仓库、重量范围、运费单价等"],
        ["编辑价目", "修改运费配置"],
        ["删除价目", "有二次确认"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("9.5  菜单管理（settings/menus）", 2)
add_note("路由：/work/settings/menus")
make_table(
    ["功能点", "说明"],
    [
        ["按角色配置可见菜单", "为每个角色勾选/取消勾选菜单项"],
        ["实时生效", "保存后该角色用户刷新页面即看到更新后的菜单"],
        ["菜单分组展示", "按「运营/财务/配置」分组显示菜单项"],
    ],
    col_widths=[5, 11],
)
gap()

add_heading("9.6  操作日志（settings/logs）", 2)
add_note("路由：/work/settings/logs")
make_table(
    ["功能点", "说明"],
    [
        ["日志列表", "展示所有用户操作记录（操作人、操作类型、目标模块、时间、IP）"],
        ["搜索/筛选", "按操作人、模块、时间范围筛选"],
        ["只读", "日志不可编辑或删除"],
    ],
    col_widths=[5, 11],
)
gap()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 十、通用功能
# ═══════════════════════════════════════════════════════════
add_heading("十、通用功能（全模块适用）", 1)
make_table(
    ["功能点", "说明"],
    [
        ["权限控制（RBAC）", "路由级鉴权：session 无效→跳转登录；角色无权限→跳转首页；菜单项按角色动态显示"],
        ["左侧导航菜单", "按角色权限动态渲染；当前页高亮；支持折叠分组"],
        ["顶部导航栏", "显示当前模块名称 + 登出按钮"],
        ["分页", "表格数据量大时分页加载"],
        ["空数据提示", "无结果时显示「暂无数据」"],
        ["加载状态", "查询中显示「查询中…」；保存中显示「保存中…」等"],
        ["Toast 通知", "操作成功/失败均有提示弹出"],
        ["弹窗关闭行为", "点取消/右上角×/弹窗外区域均可关闭；不自动保存"],
        ["图片预览", "点击图片缩略图弹出全屏预览查看器"],
        ["外链跳转", "参考链接在新标签页打开"],
        ["单位切换", "选品/询价/核价/确品/采购均支持 cmkg ↔ 英寸/英镑，仅影响显示不写入DB"],
        ["自动计算", "成本相关字段（体积重/尾程成本/temu报价等）由后端/前端自动计算，不可手动输入"],
    ],
    col_widths=[5, 11],
)
gap()
gap()

sig_p = doc.add_paragraph()
r = sig_p.add_run("整理人：____________　　整理日期：2026-04　　版本：v1.0")
_set_font(r)

out_path = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "系统功能清单.docx")
)
doc.save(out_path)
print(f"Saved: {out_path}")
