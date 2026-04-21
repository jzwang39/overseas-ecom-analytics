"""Generate 确品功能手工测试手册.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局字体 ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "���软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p

def add_para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"📌 {text}")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500
    return p

def make_table(doc, headers, rows, header_bg="1E40AF", col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, "F1F5F9")  # slate-100
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ═══════════════════════════════════════════════════════
#  封面
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("确品功能 手工测试手册")
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("ops.confirm  /  版本 1.0")
run2.font.name = "微软雅黑"
run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  一、数据准备
# ═══════════════════════════════════════════════════════
add_heading(doc, "一、数据准备", 1)
add_note(doc, "以下数据需在测试开始前通过「询价/核价」流程提前准备好，确保进入确品页有可操作的记录。")

add_heading(doc, "1.1  准备核价提交的展示记录", 2)
add_para(doc, "进入「核价」页面，新建或找一条待核价记录，填写如下字段后点击提交，使其进入确品流程：")
make_table(doc,
    ["字段", "填写值", "说明"],
    [
        ["名称", "手测-展示-001", "用于展示/单位切换测试"],
        ["所属类目", "任意已有类目", ""],
        ["产品尺寸（长/宽/高 cm）", "10 / 11 / 12", "验证 10x11x12 cm 格式"],
        ["产品重量（kg）", "1", "验证 1 kg 格式"],
        ["单套尺寸（长/宽/高 cm）", "20 / 21 / 22", ""],
        ["包裹实重（kg）", "2", ""],
    ],
    col_widths=[4.5, 5.5, 6],
)

doc.add_paragraph()
add_heading(doc, "1.2  准备操作场景记录（共 3 条）", 2)
add_para(doc, "同上，通过核价提交使以下 3 条记录均处于「待确品」状态：")
make_table(doc,
    ["名称", "用途"],
    [
        ["手测-保存-001", "测试「修改 + 保存」——状态应维持不变"],
        ["手测-提交-001", "测试「修改 + 提交」——状态变为待采购"],
        ["手测-撤回-001", "先提交至待采购，再测试「撤回」"],
    ],
    col_widths=[5, 11],
)

doc.add_paragraph()
add_heading(doc, "1.3  确认入口", 2)
add_para(doc, "打开菜单 运营 → 确品，确认页面加载成功，搜索框和状态卡片可见。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  二、页面基础展示
# ═══════════════════════════════════════════════════════
add_heading(doc, "二、页面基础展示", 1)

add_heading(doc, "TC-01  页面控件存在", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "打开确品页", "页面正常加载，无报错"],
        ["2", "查看顶部状态卡片", "同时存在「全部待确品商品」「待确品」「待采购」三个卡片"],
        ["3", "查看筛选区", "存在「商品名称」搜索框、「所属类目」下拉、「时间范围」下拉、「查询」按钮"],
        ["4", "查看表头", "存在「商品基本信息」「产品属性」「单套属性」「状态」「操作」列"],
    ],
    col_widths=[1.5, 5.5, 9],
)

doc.add_paragraph()
add_heading(doc, "TC-02  产品/包裹尺寸格式", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "搜索「手测-展示-001」并点查询", "记录出现在列表中"],
        ["2", "查看产品属性列", "显示 10x11x12 cm 和 1 kg"],
        ["3", "查看单套属性列", "显示 20x21x22 cm 和 2 kg"],
        ["4", "查看状态列", "显示「待确品」（橙色标记）"],
    ],
    col_widths=[1.5, 5.5, 9],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  三、搜索与筛选
# ═══════════════════════════════════════════════════════
add_heading(doc, "三、搜索与筛选", 1)

add_heading(doc, "TC-03  关键词搜索", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "在搜索框输入「手测-保存-001」", "—"],
        ["2", "点击「查询」", "列表只显示该记录"],
        ["3", "清空搜索框，输入不存在的名字「手测-XXXXXXXX」", "—"],
        ["4", "点击「查询」", "列表显示「暂无数据」"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()
add_heading(doc, "TC-04  类目筛选", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "从「所属类目」下拉选择准备数据时使用的类目", "—"],
        ["2", "点击「查询」", "结果只包含该类目的记录"],
        ["3", "改回「全部」，再点「查询」", "全量记录恢复"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()
add_heading(doc, "TC-05  状态卡片筛选（客户端即时过滤，无需点查询）", 2)
add_note(doc, "此筛选在客户端即时生效，切换卡片后无需再点「查询」按钮。")
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "先点「查询」加载全量数据", "列表包含待确品和待采购记录"],
        ["2", "点击「待确品」卡片", "列表仅显示待确品记录，待采购记录消失"],
        ["3", "点击「待采购」卡片", "列表仅显示待采购记录"],
        ["4", "点击「全部待确品商品」卡片", "两种状态记录均显示"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  四、修改并保存
# ═══════════════════════════════════════════════════════
add_heading(doc, "四、修改并保存（状态不变）", 1)

add_heading(doc, "TC-06  编辑弹窗字段写入", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "搜索并找到「手测-保存-001」", "—"],
        ["2", "点击该行「修改」按钮", "弹出编辑弹窗，标题包含商品名称"],
        ["3", "在「平台在售价格」Min 框填 88.8，Max 框填 99.9", "—"],
        ["4", "点击「保存」按钮", "弹窗关闭"],
        ["5", "重新搜索，再次点「修改」", "弹窗重新打开"],
        ["6", "查看「平台在售价格」", "Min = 88.8，Max = 99.9（已持久化）"],
        ["7", "查看列表中该记录状态", "仍为「待确品」（状态未变）"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  五、提交
# ═══════════════════════════════════════════════════════
add_heading(doc, "五、提交（待确品 → 待采购）", 1)

add_heading(doc, "TC-07  提交流程", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "搜索并找到「手测-提交-001」", "—"],
        ["2", "点击「修改」，打开编辑弹窗", "—"],
        ["3", "在「平台在售价格」Min 框填入 55.5", "—"],
        ["4", "点击「提交」按钮", "弹窗关闭"],
        ["5", "重新查询全部，点「待采购」状态卡片", "「手测-提交-001」出现在待采购列表"],
        ["6", "点「待确品」卡片", "「手测-提交-001」不再出现"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  六、撤回
# ═══════════════════════════════════════════════════════
add_heading(doc, "六、撤回（待采购 → 待核价）", 1)
add_note(doc, "执行 TC-08 前请确保「手测-撤回-001」已处于待采购状态（若还是待确品，先按 TC-07 步骤提交一次）。")

add_heading(doc, "TC-08  撤回必填验证", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "在待采购列表找到「手测-撤回-001」", "—"],
        ["2", "点击「撤回」按钮", "弹出撤回弹窗，标题含记录 ID"],
        ["3", "不填写理由，查看「确认撤回」按钮", "按钮为禁用（灰色），无法点击"],
        ["4", "在「撤回理由」框填写「资料不完整，需重新核价」", "—"],
        ["5", "查看「确认撤回」按钮", "按钮变为可点击（红色）"],
        ["6", "点击「确认撤回」", "弹窗关闭"],
        ["7", "在确品页搜索「手测-撤回-001」", "记录不再出现（已离开确品流程）"],
        ["8", "进入「核价」页面搜索「手测-撤回-001」", "记录出现，状态为「待核价」，撤回理由已记录"],
    ],
    col_widths=[1.5, 7, 7.5],
)

doc.add_paragraph()
add_heading(doc, "TC-09  撤回取消", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "对任意待采购记录点「撤回」", "弹窗打开"],
        ["2", "填写理由后点「取消」", "弹窗关闭，记录状态不变，仍为待采购"],
    ],
    col_widths=[1.5, 7, 7.5],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  七、单位切换
# ═══════════════════════════════════════════════════════
add_heading(doc, "七、编辑弹窗：单位切换", 1)

add_heading(doc, "TC-10  cm/kg ↔ 英寸/英镑 互换显示", 2)
add_note(doc, "使用「手测-展示-001」（产品尺寸 10/11/12 cm，重量 1 kg）。")
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "打开「手测-展示-001」的编辑弹窗", "—"],
        ["2", "查看「产品属性」区块默认单位", "长度框显示 10，「cmkg」按钮处于激活态"],
        ["3", "点击「英寸/英镑」切换按钮", "长度框变为 3.937（10cm ÷ 2.54）"],
        ["4", "点击「cmkg」切换回来", "长度框恢复 10"],
        ["5", "点「取消」关闭弹窗", "记录未被修改（仅展示换算，不写入）"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  八、历史数据（加分项）
# ═══════════════════════════════════════════════════════
add_heading(doc, "八、历史数据查看（加分项）", 1)

add_heading(doc, "TC-11  历史数据入口", 2)
make_table(doc,
    ["步骤", "操作", "预期结果"],
    [
        ["1", "在列表找到任意记录，点击「查看历史数据」按钮", "页面切换到历史数据视图"],
        ["2", "查看历史记录列表", "显示该商品的历史版本"],
        ["3", "点击「退出历史数据」", "返回正常确品列表视图"],
    ],
    col_widths=[1.5, 6.5, 8],
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  九、测试结果记录
# ═══════════════════════════════════════════════════════
add_heading(doc, "九、测试结果记录", 1)

result_rows = [
    ["TC-01", "页面控件存在", "", ""],
    ["TC-02", "尺寸格式展示", "", ""],
    ["TC-03", "关键词搜索", "", ""],
    ["TC-04", "类目筛选", "", ""],
    ["TC-05", "状态卡片筛选", "", ""],
    ["TC-06", "修改并保存", "", ""],
    ["TC-07", "提交流程", "", ""],
    ["TC-08", "撤回必填验证", "", ""],
    ["TC-09", "撤回取消", "", ""],
    ["TC-10", "单位切换", "", ""],
    ["TC-11", "历史数据（加分项）", "", ""],
]

t = make_table(doc,
    ["用例编号", "用例名称", "结果（通过/失败/跳过）", "备注"],
    result_rows,
    col_widths=[2.5, 5, 4.5, 4],
)

doc.add_paragraph()
doc.add_paragraph()

sig_p = doc.add_paragraph()
run = sig_p.add_run("测试人：____________　　测试日期：____________　　测试环境：____________")
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.size = Pt(10.5)

# ── 保存 ──────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "..", "确品功能手工测试手册.docx")
out_path = os.path.normpath(out_path)
doc.save(out_path)
print(f"Saved: {out_path}")
